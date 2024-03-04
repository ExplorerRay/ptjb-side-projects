#產生審查意見表 (給課程老師) 有些地方因檔名、檔案位置等原因需修改
import docx
import openpyxl
from docx.shared import Pt
import os

num = int(input("請輸入審查委員數:")) #要小於等於5
crs = int(input("請輸入課程數:"))
date_str = str(input("請輸入審查日期:"))

# excel 1-2
wbsec = openpyxl.load_workbook('1-2.xlsx') #resource #需修改!!
shtsec = wbsec.worksheets[0]

for v in range(crs):
    doc = docx.Document('1-3.docx') #需修改!!
    doc.styles['Normal'].font.name = "標楷體"
    doc.styles['Normal'].font.size = Pt(14)

    #加課程編號
    crs_num = shtsec.cell(row=v+2,column=1).value
    r = doc.paragraphs[3].add_run(crs_num)
    r.font.size=Pt(14)

    #更新審查日期
    doc.paragraphs[7].add_run(date_str)

    tb = doc.tables[0]
    tb.style.font.name = "標楷體" #not sure useful or not
    tb.style.font.size = Pt(14)
    #print(tb.rows[0].cells[1].text)
    for i in range(1, num+1):
        tb.rows[i].cells[0].text = str("委員"+str(i)+"\n")+shtsec.cell(row=v+2,column=i+1).value

    # print(doc.paragraphs)
    # print("paragraphs:", len(doc.paragraphs))

    # cnt=1
    # for p in doc.paragraphs:
    #     print(p.text, cnt)
    #     cnt+=1

    doc.save(str(crs_num)+'.docx')
