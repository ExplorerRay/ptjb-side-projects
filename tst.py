import docx
import openpyxl
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

num = int(input("請輸入審查委員數:"))
crs = int(input("請輸入課程數:"))

# excel 1-4 課程大表
full = openpyxl.load_workbook('1-4.xlsx') # source 需修改檔名且放在同資料夾!!
shtfull = full.worksheets[2] # 2代表終端聯盟

# excel 1-2 from google form審查意見
gf = openpyxl.load_workbook('1-2.xlsx')
shtgf = gf.worksheets[0]

for re in range(num): 
    for v in range(crs):
        # 從課程大表取資訊
        crs_num = shtfull.cell(row=v+4,column=2).value # 取課程編號
        crs_nam = shtfull.cell(row=v+4,column=8).value # 取課程名稱
        key_mod = shtfull.cell(row=v+4,column=15).value # 取重點模組
        crs_hos = shtfull.cell(row=v+4,column=7).value # 取課程主持人(教師)
        sch = shtfull.cell(row=v+4,column=3).value # 取學校
        dpt = shtfull.cell(row=v+4,column=6).value # 取系所

        # 讀docx並寫入後另存新檔
        doc = docx.Document('1-3.docx')
        tb = doc.tables[0]

        # 增加審查委員
        re_nam = shtgf.cell(row=re+2,column=3).value[0:3]
        rn = doc.paragraphs[11].add_run(re_nam) 
        rn.font.name = "標楷體"
        rn.font.size = Pt(18)
        rn._element.rPr.rFonts.set(qn('w:eastAsia'),'標楷體')

        # 綜合評分
        score = str(shtgf.cell(row=re+2,column=10+v*7).value)
        run = tb.rows[11].cells[0].paragraphs[0].add_run(score)
        run.font.name = "標楷體"
        run.font_size = Pt(14)
        run._element.rPr.rFonts.set(qn('w:eastAsia'),'標楷體')

        ans = ''
        if score=='4': ans+='■特優(4分)\n'
        else: ans+='□特優(4分)\n'
        if score=='3': ans+='■優  (3分)\n'
        else: ans+='□優  (3分)\n'
        if score=='2': ans+='■良  (2分)\n'
        else: ans+='□良  (2分)\n'
        if score=='1': ans+='■差  (1分)\n'
        else: ans+='□差  (1分)\n'
        if score=='4' or score=='3' or score=='2' or score=='1': ans+='□其他'
        else: ans+='■其他('+score+'分)'
        tb.rows[11].cells[5].text = ans

        # 設定寫入之字型及大小
        doc.styles['Normal'].font.name = "標楷體"
        doc.styles['Normal'].font.size = Pt(12)
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'),'標楷體')

        # 獲取google表單回覆
        for k in range(4): # 四項審查重點
            result = shtgf.cell(row=re+2,column=4+k+v*7).value[0] #只取答案之第一個字
            if result=='特':
                tb.rows[4+k].cells[5].text = '■特優'
                tb.rows[4+k].cells[5].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
            elif result=='優':
                tb.rows[4+k].cells[6].text = '■優'
                tb.rows[4+k].cells[6].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
            elif result=='良':
                tb.rows[4+k].cells[7].text = '■良'
                tb.rows[4+k].cells[7].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
            else:
                tb.rows[4+k].cells[8].text = '■差'
                tb.rows[4+k].cells[8].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER

        # 審查意見
        opin_fir = shtgf.cell(row=re+2,column=8+v*7).value #選取部分
        if opin_fir[0]=='無': opin_fir=''
        opin_ls = opin_fir.split(', ')
        opin_fir=''
        for o in opin_ls:
            opin_fir = opin_fir + o + '\n'
        opin_sec = shtgf.cell(row=re+2,column=9+v*7).value #打字部分
        tb.rows[9].cells[0].text = opin_fir + '\n' + opin_sec

        tb.rows[0].cells[7].text = crs_num # 增加課程編號
        tb.rows[0].cells[3].text = crs_nam # 增加課程名稱
        tb.rows[1].cells[3].text = key_mod # 增加重點模組
        tb.rows[2].cells[3].text = crs_hos # 增加課程主持人
        tb.rows[2].cells[7].text = sch + '/' + dpt # 增加服務單位

        # 以課程編號存檔
        doc.save(str(crs_num)+str(re_nam)+'.docx')


# for p in doc.paragraphs:
#     print(p.text)
#     print("=======\n")

# doc = docx.Document('1-3.docx')
# tb=doc.tables[0]
# print(tb.rows[0].cells[7].text)
# print("rows=" + str(len(tb.rows)))
# print("columns=" + str(len(tb.columns)))
