import openpyxl
import docx
from docx.shared import Pt
import os

num = int(input("請輸入審查委員數:")) #要小於等於5
cnt = int(input("請輸入課程數:"))

wbr = openpyxl.load_workbook('2-1.xlsx') #google form來的
shtr = wbr.worksheets[0] 

wb = openpyxl.load_workbook('3-1.xlsx') #課程大表要改名為3-1
sht = wb.worksheets[2] #根據聯盟決定index

for v in range(cnt):
    nm = shtr.cell(row=v+2,column=3).value
    doc = docx.Document(str(nm)+'.docx')
    doc.styles['Normal'].font.name = "標楷體"
    doc.styles['Normal'].font.size = Pt(14)

    for i in range(4, sht.max_row+1):
        if str(sht.cell(row=i,column=2).value)==nm:
            r=doc.paragraphs[4].add_run(str(sht.cell(row=i,column=3).value)+'/'+str(sht.cell(row=i,column=6).value)) #學校/系所
            doc.paragraphs[5].add_run(str(sht.cell(row=i,column=8).value)) #課程名稱
            doc.paragraphs[6].add_run(str(sht.cell(row=i,column=7).value)) #課程教師
            r.font.name="標楷體"
            r.font.size=Pt(14)
            break

    tb=doc.tables[0]
    for i in range(1, num+1):
        tb.rows[i].cells[1].text = shtr.cell(row=v+2,column=i+3).value


    doc.save(str(nm)+'審查意見回覆.docx')
