#要將此檔和1-1, 2-1放在同個資料夾 (這樣就不需要chdir)
import openpyxl
from openpyxl.styles import Alignment
import docx
from docx.shared import Pt
import os

num = int(input("請輸入審查委員數:"))
crs = int(input("請輸入課程數:"))
date_str = str(input("請輸入審查日期:"))
#print(os.getcwd())

wb_new = openpyxl.Workbook()
sht_new = wb_new.worksheets[0]

try:
    wb = openpyxl.load_workbook('1-1.xlsx') #resource
    sht = wb.worksheets[0]

    #新表預處理(尚未根據不同google form進行變化)
    sht_new['A1'] = '課程編號'
    for c in range(1, num+1):
        sht_new.cell(row=1, column=c+1).value = str(sht.cell(row=c+1,column=3).value +"之審查意見")
        #sht_new.cell(row=1, column=c+1).value = str('委員'+str(c))
    for rw in range(1, crs+1):
        sht_new.cell(row=rw+1, column=1).value = str(str(sht.cell(row=1,column=4+3*(rw-1)).value)[:6])

    #combine審查意見及其他(尚未根據不同google form進行變化，亦無法確定委員和編號之對應[若委員2比委員1先填表單，順序會錯])
    for rw in range(1, crs+1):
        for cnt in range(4):
            sht_new.cell(row = rw+1, column = 2+cnt).value = str('1.')+str(sht.cell(row = cnt+2, column = 4+3*(rw-1)).value) + str('\n') + str('2.') + str(sht.cell(row = cnt+2, column = 4+3*(rw-1)+1).value)


    for i in range(1, sht_new.max_row+1):
        for j in range(1, sht_new.max_column+1):
            print(sht_new.cell(row = i, column = j).value)
            sht_new.cell(row = i, column = j).alignment = Alignment(wrapText=True)

    wb_new.save('1-2.xlsx')
except FileNotFoundError:
    print("\nError:\n1-1.xlsx NOT found.\n請將此程式和1-1.xlsx放在同個資料夾\n")
    input("Press Any Key to close this window")

# excel 1-2
wbsec = openpyxl.load_workbook('1-2.xlsx') #resource #需修改!!
shtsec = wbsec.worksheets[0]

for v in range(crs):
    doc = docx.Document('1-3.docx') #需修改!!
    doc.styles['Normal'].font.name = "標楷體"
    doc.styles['Normal'].font.size = Pt(14)

    #加課程編號
    crs_num = shtsec.cell(row=v+2,column=1).value
    r = doc.paragraphs[3].add_run(crs_num)
    r.font.size=Pt(14)

    #更新審查日期
    doc.paragraphs[7].add_run(date_str)

    tb = doc.tables[0]
    tb.style.font.name = "標楷體" #not sure useful or not
    tb.style.font.size = Pt(14)
    #print(tb.rows[0].cells[1].text)
    for i in range(1, num+1):
        tb.rows[i].cells[0].text = str("委員"+str(i)+"\n")+shtsec.cell(row=v+2,column=i+1).value
    
    doc.save(str(crs_num)+'.docx')

input("Press Any Key to close this window")
